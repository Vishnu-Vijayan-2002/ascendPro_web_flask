from flask import Blueprint, render_template, request, redirect, send_file, session, flash

import sqlite3
from config import DATABASE
from nlp.skill_extractor import extract_skills
from nlp.ats_scorer import calculate_ats_score
from werkzeug.utils import secure_filename
import PyPDF2
import docx
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io
import os
from datetime import datetime

applicant_bp = Blueprint("applicant", __name__)

ALLOWED_EXTENSIONS = {'pdf', 'docx', 'txt'}

# Absolute path — fixes FileNotFoundError on Windows
BASE_DIR      = os.path.dirname(os.path.abspath(__file__))
UPLOAD_FOLDER = os.path.join(BASE_DIR, '..', 'uploads')


# ══════════════════════════════════════════════════════════════════
# FILE HELPERS
# ══════════════════════════════════════════════════════════════════

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def extract_text_from_pdf(file):
    try:
        pdf_reader = PyPDF2.PdfReader(file)
        text = ""
        for page in pdf_reader.pages:
            text += page.extract_text() + "\n"
        return text.strip()
    except Exception as e:
        print(f"Error extracting PDF: {e}")
        return ""

def extract_text_from_docx(file):
    try:
        doc  = docx.Document(file)
        text = "\n".join([p.text for p in doc.paragraphs])
        return text.strip()
    except Exception as e:
        print(f"Error extracting DOCX: {e}")
        return ""

def extract_text_from_txt(file):
    try:
        return file.read().decode('utf-8').strip()
    except Exception as e:
        print(f"Error extracting TXT: {e}")
        return ""


# ══════════════════════════════════════════════════════════════════
# ROUTES
# ══════════════════════════════════════════════════════════════════

# ── Dashboard ──────────────────────────────────────────────────────
@applicant_bp.route("/dashboard")
def dashboard():

    user_id = session.get("user_id")

    total_applications = 0
    shortlisted = 0
    accepted = 0

    if user_id:
        conn = sqlite3.connect(DATABASE)
        cur = conn.cursor()

        # Total Applications
        cur.execute("SELECT COUNT(*) FROM applications WHERE user_id = ?", (user_id,))
        total_applications = cur.fetchone()[0]

        # Shortlisted
        cur.execute("""
            SELECT COUNT(*) FROM applications 
            WHERE user_id = ? AND status = 'Shortlisted'
        """, (user_id,))
        shortlisted = cur.fetchone()[0]

        # Accepted
        cur.execute("""
            SELECT COUNT(*) FROM applications 
            WHERE user_id = ? AND status = 'Accepted'
        """, (user_id,))
        accepted = cur.fetchone()[0]

        conn.close()

    return render_template(
        "applicant/dashboard.html",
        total_applications=total_applications,
        shortlisted=shortlisted,
        accepted=accepted
    )



# ── Upload Resume ──────────────────────────────────────────────────
@applicant_bp.route("/upload-resume", methods=["GET", "POST"])
def upload_resume():
    message   = ""
    skills    = []
    ats_score = None

    if request.method == "POST":
        user_id = request.form.get("user_id")
        content = ""

        if 'resume' not in request.files:
            message = "⚠️ No file uploaded"
            return render_template("applicant/upload_resume.html",
                                   message=message, skills=skills, ats_score=ats_score)

        file = request.files['resume']

        if file.filename == '':
            message = "⚠️ No file selected"
            return render_template("applicant/upload_resume.html",
                                   message=message, skills=skills, ats_score=ats_score)

        if file and allowed_file(file.filename):
            filename = secure_filename(file.filename)
            file_ext = filename.rsplit('.', 1)[1].lower()
            file.stream.seek(0)

            if file_ext == 'pdf':
                content = extract_text_from_pdf(file)
            elif file_ext == 'docx':
                content = extract_text_from_docx(file)
            elif file_ext == 'txt':
                content = extract_text_from_txt(file)

            if not content or len(content.strip()) < 10:
                message = "⚠️ Could not extract text. Please try another format."
                return render_template("applicant/upload_resume.html",
                                       message=message, skills=skills, ats_score=ats_score)
        else:
            message = "⚠️ Invalid format. Upload PDF, DOCX, or TXT."
            return render_template("applicant/upload_resume.html",
                                   message=message, skills=skills, ats_score=ats_score)

        skills    = extract_skills(content)
        ats_score = calculate_ats_score(content)

        try:
            conn = sqlite3.connect(DATABASE)
            cur  = conn.cursor()
            cur.execute(
                "INSERT INTO resumes (user_id, content, ats_score) VALUES (?, ?, ?)",
                (user_id, content, ats_score)
            )
            conn.commit()
            conn.close()
            message = "✅ Resume uploaded successfully!"
        except Exception as e:
            message = f"⚠️ Database error: {str(e)}"
            print(f"Database error: {e}")

    return render_template("applicant/upload_resume.html",
                           message=message, skills=skills, ats_score=ats_score)


# ── Build Resume ───────────────────────────────────────────────────
@applicant_bp.route("/build-resume", methods=["GET", "POST"])
def build_resume():
    if request.method == "POST":
        user_id     = request.form.get("user_id")
        resume_data = collect_resume_data(request.form)
        content     = generate_professional_resume(resume_data)
        ats_score   = calculate_ats_score(content)
        skills      = extract_skills(content)

        try:
            conn = sqlite3.connect(DATABASE)
            cur  = conn.cursor()
            cur.execute(
                "INSERT INTO resumes (user_id, content, ats_score) VALUES (?, ?, ?)",
                (user_id, content, ats_score)
            )
            resume_id = cur.lastrowid
            conn.commit()
            conn.close()

            return render_template(
                "applicant/resume_preview.html",
                resume_id=resume_id, user_id=user_id,
                content=content, ats_score=ats_score,
                skills=skills, resume_data=resume_data,
                message="✅ Resume created successfully!"
            )
        except Exception as e:
            print(f"Database error: {e}")
            return render_template("applicant/build_resume.html",
                                   error=f"⚠️ Error saving resume: {str(e)}")

    return render_template("applicant/build_resume.html")


# ── Edit Resume ────────────────────────────────────────────────────
@applicant_bp.route("/edit-resume/<int:resume_id>", methods=["GET", "POST"])
def edit_resume(resume_id):
    if request.method == "POST":
        user_id     = request.form.get("user_id")
        resume_data = collect_resume_data(request.form)
        content     = generate_professional_resume(resume_data)
        ats_score   = calculate_ats_score(content)
        skills      = extract_skills(content)

        try:
            conn = sqlite3.connect(DATABASE)
            cur  = conn.cursor()
            cur.execute(
                "UPDATE resumes SET content = ?, ats_score = ? WHERE id = ?",
                (content, ats_score, resume_id)
            )
            conn.commit()
            conn.close()

            return render_template(
                "applicant/resume_preview.html",
                resume_id=resume_id, user_id=user_id,
                content=content, ats_score=ats_score,
                skills=skills, resume_data=resume_data,
                message="✅ Resume updated successfully!"
            )
        except Exception as e:
            print(f"Database error: {e}")
            return render_template("applicant/build_resume.html",
                                   error=f"⚠️ Error updating resume: {str(e)}")

    try:
        conn = sqlite3.connect(DATABASE)
        cur  = conn.cursor()
        cur.execute("SELECT content, user_id FROM resumes WHERE id = ?", (resume_id,))
        result = cur.fetchone()
        conn.close()
        if result:
            return render_template("applicant/build_resume.html")
        return redirect("/applicant/dashboard")
    except Exception as e:
        print(f"Error: {e}")
        return redirect("/applicant/dashboard")


# ── Download Resume ────────────────────────────────────────────────
@applicant_bp.route("/download-resume/<int:resume_id>")
def download_resume(resume_id):
    try:
        conn = sqlite3.connect(DATABASE)
        cur  = conn.cursor()
        cur.execute("SELECT content, user_id FROM resumes WHERE id = ?", (resume_id,))
        result = cur.fetchone()
        conn.close()

        if not result:
            return "Resume not found", 404

        content, user_id = result
        doc         = create_professional_docx(content)
        file_stream = io.BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)

        filename = f"Resume_{user_id}_{datetime.now().strftime('%Y%m%d')}.docx"
        return send_file(
            file_stream,
            as_attachment=True,
            download_name=filename,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
    except Exception as e:
        print(f"Download error: {e}")
        return f"Error: {str(e)}", 500


# ── View Jobs ──────────────────────────────────────────────────────
@applicant_bp.route("/jobs")
def jobs():

    user_id = session.get("user_id")

    conn = sqlite3.connect(DATABASE)
    cur  = conn.cursor()

    # Get all jobs
    cur.execute("SELECT * FROM jobs")
    jobs = cur.fetchall()

    applied_jobs = {}

    # If user logged in → fetch their applications
    if user_id:
        cur.execute("""
            SELECT job_id, status 
            FROM applications 
            WHERE user_id = ?
        """, (user_id,))

        applications = cur.fetchall()

        # Convert to dictionary
        # { job_id: status }
        applied_jobs = {app[0]: app[1] for app in applications}

    conn.close()

    return render_template(
        "applicant/job_list.html",
        jobs=jobs,
        applied_jobs=applied_jobs
    )


# ── Apply for Job ──────────────────────────────────────────────────
@applicant_bp.route("/apply/<int:job_id>", methods=["GET", "POST"])
def apply(job_id):

    print("\n=== APPLY DEBUG ===")
    print("SESSION      :", dict(session))
    print("USER_ID      :", session.get("user_id"))
    print("METHOD       :", request.method)
    print("===================\n")

    user_id = session.get("user_id")

    if not user_id:
        flash("⚠️ Please login first.", "danger")
        return redirect("/auth/login")

    # Fetch job details
    conn = sqlite3.connect(DATABASE)
    cur  = conn.cursor()
    cur.execute("SELECT * FROM jobs WHERE id = ?", (job_id,))
    job = cur.fetchone()

    if not job:
        conn.close()
        return redirect("/applicant/jobs")

    # Check if already applied (for GET & POST both)
    cur.execute(
        "SELECT id FROM applications WHERE user_id=? AND job_id=?",
        (user_id, job_id)
    )
    already = cur.fetchone()

    if request.method == "POST":

        if already:
            conn.close()
            flash("⚠️ You have already applied for this job.", "warning")
            return redirect(f"/applicant/apply/{job_id}")

        full_name    = request.form.get("full_name", "").strip()
        email        = request.form.get("email", "").strip()
        phone        = request.form.get("phone", "").strip()
        education    = request.form.get("education", "").strip()
        cover_letter = request.form.get("cover_letter", "").strip()
        resume       = request.files.get("resume")

        if not full_name or not email or not phone:
            conn.close()
            flash("⚠️ Full name, email and phone are required.", "danger")
            return redirect(f"/applicant/apply/{job_id}")

        resume_filename = None

        if resume and resume.filename != "":
            if not allowed_file(resume.filename):
                conn.close()
                flash("⚠️ Invalid file. Upload PDF, DOCX, or TXT only.", "danger")
                return redirect(f"/applicant/apply/{job_id}")

            resume_filename = f"{user_id}_{job_id}_{secure_filename(resume.filename)}"
            os.makedirs(UPLOAD_FOLDER, exist_ok=True)
            resume.save(os.path.join(UPLOAD_FOLDER, resume_filename))
            print("✅ Resume saved:", resume_filename)

        try:
            cur.execute("""
                INSERT INTO applications
                  (user_id, job_id, status, full_name, email, phone,
                   education, cover_letter, resume_filename)
                VALUES (?, ?, 'Applied', ?, ?, ?, ?, ?, ?)
            """, (
                user_id, job_id, full_name, email, phone,
                education, cover_letter, resume_filename
            ))

            cur.execute(
                "INSERT INTO notifications (user_id, message) VALUES (?, ?)",
                (user_id, f"📩 You applied for: {job[1]}")
            )

            conn.commit()
            conn.close()

            flash("✅ Application submitted successfully!", "success")
            return redirect("/applicant/jobs")

        except Exception as e:
            conn.close()
            flash(f"⚠️ Database error: {str(e)}", "danger")
            return redirect(f"/applicant/apply/{job_id}")

    conn.close()

    return render_template(
        "applicant/apply_job.html",
        job=job,
        job_id=job_id,
        already_applied=bool(already)
    )

# ══════════════════════════════════════════════════════════════════
# RESUME HELPERS
# ══════════════════════════════════════════════════════════════════

def collect_resume_data(form):
    data = {
        'full_name':        form.get('full_name'),
        'location':         form.get('location'),
        'phone':            form.get('phone'),
        'email':            form.get('email'),
        'linkedin':         form.get('linkedin', ''),
        'github':           form.get('github', ''),
        'portfolio':        form.get('portfolio', ''),
        'summary':          form.get('summary'),
        'skills_languages': form.get('skills_languages', ''),
        'skills_frontend':  form.get('skills_frontend', ''),
        'skills_backend':   form.get('skills_backend', ''),
        'skills_database':  form.get('skills_database', ''),
        'skills_tools':     form.get('skills_tools', ''),
        'educations':       [],
        'projects':         [],
        'experiences':      [],
        'certifications':   []
    }

    edu_count = int(form.get('education_count', 0))
    for i in range(edu_count):
        data['educations'].append({
            'degree':      form.get(f'edu_degree_{i}'),
            'institution': form.get(f'edu_institution_{i}'),
            'duration':    form.get(f'edu_duration_{i}'),
            'details':     form.get(f'edu_details_{i}', '')
        })

    proj_count = int(form.get('projects_count', 0))
    for i in range(proj_count):
        data['projects'].append({
            'name':        form.get(f'project_name_{i}'),
            'date':        form.get(f'project_date_{i}', ''),
            'github':      form.get(f'project_github_{i}', ''),
            'live':        form.get(f'project_live_{i}', ''),
            'description': form.get(f'project_description_{i}'),
            'tech':        form.get(f'project_tech_{i}')
        })

    exp_count = int(form.get('experience_count', 0))
    for i in range(exp_count):
        data['experiences'].append({
            'title':       form.get(f'exp_title_{i}'),
            'company':     form.get(f'exp_company_{i}'),
            'location':    form.get(f'exp_location_{i}', ''),
            'duration':    form.get(f'exp_duration_{i}'),
            'description': form.get(f'exp_description_{i}')
        })

    cert_count = int(form.get('certifications_count', 0))
    for i in range(cert_count):
        data['certifications'].append({
            'name':        form.get(f'cert_name_{i}'),
            'org':         form.get(f'cert_org_{i}'),
            'date':        form.get(f'cert_date_{i}', ''),
            'credential':  form.get(f'cert_credential_{i}', ''),
            'description': form.get(f'cert_description_{i}', '')
        })

    return data


def generate_professional_resume(data):
    content = f"{data['full_name']}\n"

    links = []
    if data['linkedin']:  links.append(f"LinkedIn: {data['linkedin']}")
    if data['github']:    links.append(f"GitHub: {data['github']}")
    if data['portfolio']: links.append(f"Portfolio: {data['portfolio']}")

    content += f"{data['location']} | {data['phone']} | {data['email']}"
    if links:
        content += " | " + " | ".join(links)
    content += "\n\n"

    if data['summary']:
        content += "Professional Summary\n"
        content += f"{data['summary']}\n\n"

    if data['educations']:
        content += "Education\n"
        for edu in data['educations']:
            if edu['degree'] and edu['institution']:
                content += f"{edu['degree']}\n"
                content += f"{edu['institution']}"
                if edu['duration']: content += f" | {edu['duration']}"
                content += "\n"
                if edu['details']: content += f"{edu['details']}\n"
                content += "\n"

    content += "Technical Skills\n"
    if data['skills_languages']: content += f"Languages: {data['skills_languages']}\n"
    if data['skills_frontend']:  content += f"Frontend: {data['skills_frontend']}\n"
    if data['skills_backend']:   content += f"Backend: {data['skills_backend']}\n"
    if data['skills_database']:  content += f"Database: {data['skills_database']}\n"
    if data['skills_tools']:     content += f"Tools: {data['skills_tools']}\n"
    content += "\n"

    if data['projects']:
        content += "Projects\n"
        for proj in data['projects']:
            if proj['name']:
                proj_links = []
                if proj['github']: proj_links.append(f"GitHub: {proj['github']}")
                if proj['live']:   proj_links.append(f"Live: {proj['live']}")
                content += f"{proj['name']}"
                if proj_links: content += f" | {' | '.join(proj_links)}"
                if proj['date']: content += f" | {proj['date']}"
                content += "\n"
                if proj['description']: content += f"{proj['description']}\n"
                if proj['tech']:        content += f"Tech Stack: {proj['tech']}\n"
                content += "\n"

    if data['experiences']:
        content += "Professional Experience\n"
        for exp in data['experiences']:
            if exp['title'] and exp['company']:
                content += f"{exp['title']} | {exp['company']}"
                if exp['location']: content += f" | {exp['location']}"
                content += "\n"
                if exp['duration']:    content += f"{exp['duration']}\n"
                if exp['description']: content += f"{exp['description']}\n"
                content += "\n"

    if data['certifications']:
        content += "Certifications\n"
        for cert in data['certifications']:
            if cert['name'] and cert['org']:
                content += f"{cert['name']} – {cert['org']}"
                if cert['date']: content += f" | {cert['date']}"
                content += "\n"
                if cert['description']:  content += f"{cert['description']}\n"
                if cert['credential']:   content += f"Credential: {cert['credential']}\n"
                content += "\n"

    return content


def create_professional_docx(content):
    doc = Document()

    for section in doc.sections:
        section.top_margin    = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin   = Inches(0.75)
        section.right_margin  = Inches(0.75)

    SECTION_HEADERS = [
        'Professional Summary', 'Education', 'Technical Skills',
        'Projects', 'Professional Experience', 'Certifications'
    ]

    for i, line in enumerate(content.split('\n')):
        if not line.strip():
            continue

        if i == 0:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(line)
            run.font.size = Pt(18)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0, 0, 0)

        elif i == 1:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(line)
            run.font.size = Pt(10)

        elif line in SECTION_HEADERS:
            p = doc.add_paragraph()
            run = p.add_run(line.upper())
            run.font.size = Pt(12)
            run.font.bold = True
            run.font.color.rgb = RGBColor(31, 78, 121)
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after  = Pt(6)

        elif line.strip().startswith('•') or line.strip().startswith('-'):
            p = doc.add_paragraph(line.strip(), style='List Bullet')
            p.paragraph_format.left_indent = Inches(0.25)
            if p.runs:
                p.runs[0].font.size = Pt(10)

        else:
            p   = doc.add_paragraph(line)
            run = p.runs[0] if p.runs else p.add_run(line)
            run.font.size = Pt(10)

    return doc